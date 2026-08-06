"""python-docx를 이용한 Word 문서 생성 스크립트

- 제목: 40px(-> 30pt), 본문: 20px(-> 15pt), 폰트: 맑은 고딕
- 메인 컬러: 실행 시 인자로 원하는 색상(hex)을 지정하면 제목/표 헤더 등에 반영됨
  (Word는 폰트 크기를 pt 단위로 다뤄서, 웹 기준 96dpi 환산식 1px = 0.75pt 로 변환)
"""
import argparse

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"

PX_TO_PT = 0.75
TITLE_SIZE = Pt(40 * PX_TO_PT)     # 30pt
SUBHEAD_SIZE = Pt(28 * PX_TO_PT)   # 21pt (소제목용으로 추가)
BODY_SIZE = Pt(20 * PX_TO_PT)      # 15pt

DEFAULT_MAIN_COLOR = "2E74B5"


def _set_east_asian_font(element, font_name=FONT_NAME):
    """맑은 고딕 같은 한글 폰트는 rFonts의 eastAsia 속성도 지정해야
    Word에서 실제로 적용된다 (ascii 속성만으로는 한글에 반영되지 않음)."""
    rPr = element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color_hex=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    _set_east_asian_font(run._element)


def set_default_style(doc):
    """Normal 스타일 기본값을 지정해 새로 추가되는 문단도 동일한 폰트/크기를 따르게 한다."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    _set_east_asian_font(style.element)


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def create_document(
    output_path: str = "output.docx",
    main_color: str = DEFAULT_MAIN_COLOR,
    title_text: str = "문서 제목",
    author: str = None,
):
    main_color = main_color.lstrip("#").upper()

    doc = Document()

    # 문서 메타데이터
    doc.core_properties.title = title_text
    if author:
        doc.core_properties.author = author

    # 여백
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    set_default_style(doc)

    # 제목 (40px -> 30pt, 메인 컬러 적용)
    title_para = doc.add_heading("", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title_para.add_run(title_text), size=TITLE_SIZE, bold=True, color_hex=main_color)

    # 소제목 1
    h1 = doc.add_heading("", level=1)
    set_run_font(h1.add_run("1. 개요"), size=SUBHEAD_SIZE, bold=True, color_hex=main_color)

    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "이 문단은 python-docx로 생성된 예제 본문입니다. "
            "제목은 40px(30pt), 본문은 20px(15pt), 폰트는 맑은 고딕으로 통일되어 있습니다."
        )
    )

    # 소제목 2 - 목록
    h2 = doc.add_heading("", level=1)
    set_run_font(h2.add_run("2. 목록 예시"), size=SUBHEAD_SIZE, bold=True, color_hex=main_color)

    for item in ["첫 번째 항목", "두 번째 항목", "세 번째 항목"]:
        bullet = doc.add_paragraph(style="List Bullet")
        set_run_font(bullet.add_run(item))

    # 소제목 3 - 표
    h3 = doc.add_heading("", level=1)
    set_run_font(h3.add_run("3. 표 예시"), size=SUBHEAD_SIZE, bold=True, color_hex=main_color)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["이름", "부서", "비고"]
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, headers):
        set_cell_background(cell, main_color)
        set_run_font(cell.paragraphs[0].add_run(text), bold=True, color_hex="FFFFFF")

    rows_data = [
        ("홍길동", "개발팀", ""),
        ("김철수", "디자인팀", ""),
    ]
    for name, dept, note in rows_data:
        row_cells = table.add_row().cells
        for cell, text in zip(row_cells, (name, dept, note)):
            set_run_font(cell.paragraphs[0].add_run(text))

    # 소제목 4 - 서식
    h4 = doc.add_heading("", level=1)
    set_run_font(h4.add_run("4. 서식 예시"), size=SUBHEAD_SIZE, bold=True, color_hex=main_color)

    p2 = doc.add_paragraph()
    set_run_font(p2.add_run("굵은 글씨"), bold=True)
    set_run_font(p2.add_run(" 와 "))
    set_run_font(p2.add_run("기울임 글씨"), italic=True)
    set_run_font(p2.add_run(" 를 함께 사용할 수 있습니다."))
    set_run_font(p2.add_run(" 강조 색상 텍스트"), color_hex=main_color, bold=True)

    # 페이지 하단 쪽번호
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = footer_para.add_run()
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    set_run_font(run, size=Pt(9))

    doc.save(output_path)
    print(f"문서가 생성되었습니다: {output_path}")


def parse_args():
    parser = argparse.ArgumentParser(description="python-docx 문서 생성 스크립트")
    parser.add_argument("-o", "--output", default="output.docx", help="저장할 파일 경로")
    parser.add_argument("-c", "--color", default=DEFAULT_MAIN_COLOR, help="메인 컬러 (예: 2E74B5 또는 #2E74B5)")
    parser.add_argument("-t", "--title", default="문서 제목", help="문서 제목")
    parser.add_argument("-a", "--author", default=None, help="작성자")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    create_document(
        output_path=args.output,
        main_color=args.color,
        title_text=args.title,
        author=args.author,
    )
