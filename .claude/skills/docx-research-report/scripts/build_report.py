"""조사 내용을 담은 JSON을 입력받아 표지 - 목차 - 본문 - 참고자료 구조의
Word(.docx) 조사 보고서를 생성한다.

내용(콘텐츠)은 스크립트가 만들지 않는다. 리서치와 글쓰기는 Claude가 하고,
이 스크립트는 정해진 스키마의 JSON을 받아 서식(제목 40px/30pt, 본문 20px/15pt,
맑은 고딕, 메인 컬러)을 일관되게 입히는 역할만 한다.
JSON 스키마는 ../references/content_schema.md 참고.

사용법:
    python build_report.py --input content.json --output report.docx [--color 2E74B5]
"""
import argparse
import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"

PX_TO_PT = 0.75
TITLE_SIZE = Pt(40 * PX_TO_PT)       # 30pt - 보고서 제목
SUBTITLE_SIZE = Pt(24 * PX_TO_PT)    # 18pt - 부제
SECTION_SIZE = Pt(28 * PX_TO_PT)     # 21pt - "1. 조사 개요" 급 섹션 제목
SUBSECTION_SIZE = Pt(24 * PX_TO_PT)  # 18pt - "4.1 ..." 급 하위 섹션 제목
BODY_SIZE = Pt(20 * PX_TO_PT)        # 15pt - 본문/목차/참고자료

DEFAULT_MAIN_COLOR = "2E74B5"


# ---------- 폰트/서식 유틸 ----------

def _set_east_asian_font(element, font_name=FONT_NAME):
    """맑은 고딕 같은 한글 폰트는 rFonts의 eastAsia 속성도 지정해야
    Word에서 실제로 적용된다 (ascii 속성만으로는 한글에 반영되지 않음)."""
    rPr = element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, size=BODY_SIZE, bold=False, italic=False, color_hex=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    _set_east_asian_font(run._element)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    _set_east_asian_font(style.element)


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_hyperlink(paragraph, url, text, color_hex="0563C1"):
    """python-docx는 하이퍼링크를 기본 지원하지 않아 관계(relationship)를
    직접 추가하는 표준 레시피를 사용한다. 참고자료 출처 URL을 클릭 가능하게 만든다."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(BODY_SIZE.pt * 2)))
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_break(doc):
    doc.add_page_break()


# ---------- 문서 빌딩 블록 ----------

def build_title_page(doc, content, main_color):
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title_para.add_run(content["title"]), size=TITLE_SIZE, bold=True, color_hex=main_color)

    if content.get("subtitle"):
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(sub_para.add_run(f"- {content['subtitle']} -"), size=SUBTITLE_SIZE, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    meta_lines = []
    if content.get("date"):
        meta_lines.append(f"작성일: {content['date']}")
    if content.get("author"):
        meta_lines.append(f"작성자: {content['author']}")
    if content.get("doc_type"):
        meta_lines.append(f"문서 구분: {content['doc_type']}")

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(line))

    add_page_break(doc)


def _collect_toc_entries(sections):
    entries = []
    for section in sections:
        entries.append((1, section["title"]))
        for sub in section.get("subsections", []):
            entries.append((2, sub["title"]))
    return entries


def build_toc(doc, content, main_color):
    heading = doc.add_paragraph()
    set_run_font(heading.add_run("목차"), size=SECTION_SIZE, bold=True, color_hex=main_color)

    entries = _collect_toc_entries(content["sections"])
    if content.get("references"):
        entries.append((1, "참고 자료(출처)"))

    for level, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8 * level)
        set_run_font(p.add_run(title))

    add_page_break(doc)


def build_blocks(doc, blocks, main_color):
    for block in blocks:
        btype = block.get("type", "paragraph")
        if btype == "paragraph":
            p = doc.add_paragraph()
            set_run_font(p.add_run(block["text"]))
        elif btype == "bullets":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                set_run_font(p.add_run(item))
        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for cell, text in zip(hdr_cells, headers):
                set_cell_background(cell, main_color)
                set_run_font(cell.paragraphs[0].add_run(text), bold=True, color_hex="FFFFFF")
            for row in rows:
                row_cells = table.add_row().cells
                for cell, text in zip(row_cells, row):
                    set_run_font(cell.paragraphs[0].add_run(str(text)))
            doc.add_paragraph()
        else:
            raise ValueError(f"알 수 없는 block type: {btype}")


def build_section(doc, section, main_color):
    heading = doc.add_paragraph()
    set_run_font(heading.add_run(section["title"]), size=SECTION_SIZE, bold=True, color_hex=main_color)

    build_blocks(doc, section.get("blocks", []), main_color)

    for sub in section.get("subsections", []):
        sub_heading = doc.add_paragraph()
        set_run_font(sub_heading.add_run(sub["title"]), size=SUBSECTION_SIZE, bold=True)
        build_blocks(doc, sub.get("blocks", []), main_color)


def build_references(doc, references, main_color):
    heading = doc.add_paragraph()
    set_run_font(heading.add_run("참고 자료(출처)"), size=SECTION_SIZE, bold=True, color_hex=main_color)

    for i, ref in enumerate(references, start=1):
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{i}. {ref['text']}"))
        if ref.get("url"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.6)
            add_hyperlink(p2, ref["url"], ref["url"])

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(closing.add_run("- 끝 -"))


# ---------- 메인 ----------

def build_report(content: dict, output_path: str, main_color: str = None):
    main_color = (main_color or content.get("main_color") or DEFAULT_MAIN_COLOR).lstrip("#").upper()

    doc = Document()
    doc.core_properties.title = content["title"]
    if content.get("author"):
        doc.core_properties.author = content["author"]

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    set_default_style(doc)

    build_title_page(doc, content, main_color)
    build_toc(doc, content, main_color)

    for sec in content["sections"]:
        build_section(doc, sec, main_color)

    if content.get("references"):
        build_references(doc, content["references"], main_color)

    doc.save(output_path)
    print(f"문서가 생성되었습니다: {output_path}")


def parse_args():
    parser = argparse.ArgumentParser(description="조사 보고서 JSON -> docx 변환 스크립트")
    parser.add_argument("-i", "--input", required=True, help="콘텐츠 JSON 파일 경로")
    parser.add_argument("-o", "--output", required=True, help="저장할 docx 파일 경로")
    parser.add_argument("-c", "--color", default=None, help="메인 컬러 hex (JSON의 main_color보다 우선)")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    with open(args.input, "r", encoding="utf-8") as f:
        content = json.load(f)
    build_report(content, args.output, args.color)
